# How to comment in PyCharm:  select and CTRL+/
# How to prepare input data:
        # Filter PCAP in Wireshark by (cflow && !icmp):
                # File > Export Packet Dissections > As CSV,
                # File > Export Packet Dissections > As JSON
        # Excel > Data tab > Get data > CSV > set Data Type Detection to "Do not detect data type" > save as .xlsx

# No.	Time	Source	Destination	Protocol	Length	Source Port	Destination Port	L7 Application Name	Observation Domain Id	DstPort	DstAddr	Info
# 1	52:38.7	10.103.0.116	10.103.0.84	CFLOW	1430			https,https,https,https,https,https,https,https,https,https,https,https,https	200	443,525,004,434,915,000,000,000,000,000,000,000,000,000,000,000,000	157.240.25.63,5.37.163.15,52.221.39.79,85.154.2.101,149.202.221.211,5.36.52.122,69.195.186.128,82.178.219.180,213.202.0.161,37.41.133.105,216.58.208.74,85.154.183.126,23.9.179.248	IPFIX flow (1388 bytes) Obs-Domain-ID=  200 [Data:256]
# 2	52:38.7	10.103.0.117	10.103.0.84	CFLOW	1354			chat,amazon-kindle,amazon-kindle,too_early,facebook,too_early,too_early,facebook,https,https,too_early,too_early	200	551,994,434,868,334,000,000,000,000,000,000,000,000,000,000,000	5.36.181.164,52.95.121.236,5.36.59.161,5.37.174.236,213.202.0.32,91.108.56.108,85.154.175.179,82.178.85.114,172.217.18.138,37.41.187.88,161.117.68.190,172.217.169.226	IPFIX flow (1312 bytes) Obs-Domain-ID=  200 [Data:256]

import time
import docx
import xlrd
import json

def msecdiff(flow1, flow2):      # refer to Why#1   c2p1med01.NetFlow_first130NF_wTemplate_JSON.json
    """used to compute millisecond level difference between two IPFIX flowStartMilliseconds values in format Jul 10, 2019 11:54:15.210000"""
    return abs((int((flow1.split("."))[1]) / 10 ** 3 + int(time.mktime(time.strptime(flow1, pattern)))*1000) - (int((flow2.split("."))[1]) / 10 ** 3 + int(time.mktime(time.strptime(flow2, pattern)))*1000))  # diff in milliseconds

### fixed vars ###
commonflowsdoc = docx.Document()
commonflowsdocAbsPath = input("Please input absolute path to the file where matching L7 App & dst.port & dst.address flows and their flowStart difference in msec should be stored (like C:\\Users\\boburciu\\Downloads\\same.docx ):")
pattern = '%b %d, %Y %H:%M:%S.%f'
flowstart1 = ""
flowstart2 = ""
dictframel7app = {}
dictframeodid = {}
dictframedstport = {}
dictframedstaddr = {}
#dictframestarttime = {}
diff = 0

### changing vars - c2p1med01.NetFlow PCAP's first 130 IPFIX frames ###
with open('C:\\Users\\boburciu\\Desktop\\c2p1med01.NetFlow_first130NF_wTemplate_JSON.json') as f:
	data = json.load(f)
sh = xlrd.open_workbook("C:\\Users\\boburciu\\Desktop\\c2p1med01.NetFlow_first130NF_wTemplate.xlsx").sheet_by_index(0)
zz = 130
framenr = sh.col_values(0, start_rowx=2, end_rowx=zz)
l7app = sh.col_values(11, start_rowx=2, end_rowx=zz)
odid = sh.col_values(12, start_rowx=2, end_rowx=zz)
dstport = sh.col_values(13, start_rowx=2, end_rowx=zz)
dstaddr = sh.col_values(14, start_rowx=2, end_rowx=zz)

# sh = xlrd.open_workbook("C:\\Users\\boburciu\\Desktop\\NF_L7app_Dport_Daddr.xlsx").sheet_by_index(0) - Working

# ### changing vars - Export AppStack for IXIA 2.pcapng ###
# with open('C:\\Users\\boburciu\\Desktop\\Export AppStack for IXIA 2_JSON.json') as f:
# 	data = json.load(f)
# sh = xlrd.open_workbook("C:\\Users\\boburciu\\Desktop\\Export AppStack for IXIA 2_CSV.xlsx").sheet_by_index(0)
# framenr = sh.col_values(0, start_rowx=2)
# l7app = sh.col_values(11, start_rowx=2)
# odid = sh.col_values(14, start_rowx=2)
# dstport = sh.col_values(12, start_rowx=2)
# dstaddr = sh.col_values(13, start_rowx=2)
# starttime = sh.col_values(15, start_rowx=2)   # used to read the flowStartMilliseconds values from a CSV, but Wireshark export to CSV was found to be limited to 255 chas, so exporting as JSON and parsing that was implemented

for k in range(len(l7app)):
    dictframel7app[k] = l7app[k].split(",")       # refer to Why#3
    # dictframel7app = {0: ['https', 'https', 'https', 'https', 'https', 'https', 'https', 'https', 'https', 'https', 'https', 'https', 'https'],
    #                   1: ['chat', 'amazon-kindle', 'amazon-kindle', 'too_early', 'facebook', 'too_early', 'too_early', 'facebook', 'https', 'https', 'too_early', 'too_early'],
    #                   2: ['too_early', 'too_early', 'too_early', 'too_early', 'too_early', 'smtp', 'smtp', 'office365', 'office365', 'https', 'https', 'facebook'],
    #                   3: ['https', 'facebook', 'facebook', 'https', 'https', 'https', 'https', 'accuweather.com', 'accuweather.com'],
    #                   4: ['facebook', 'https', 'https', 'too_early', 'facebook', 'facebook', 'too_early', 'too_early', 'https', 'https', 'hpvirtgrp', 'hpvirtgrp', 'https'],
    #                   5: ['https', 'https', 'https', 'documentum_s', 'documentum_s', 'facebook', 'facebook', 'https', 'too_early', 'https', 'too_early', 'too_early'],
    #                   6: ['https', 'https', 'https', 'https', 'youtube', 'youtube', 'chat', 'chat', 'https', 'https', 'http', 'http', 'zabbix-agent', 'zabbix-agent'],
    #                   7: ['too_early', 'https', 'https', 'https', 'https', 'too_early', 'https', 'https', 'https', 'https', 'https'],
    #                   8: ['youtube', 'youtube', 'hwwlcs.com', 'hwwlcs.com', 'icloudweb', 'icloudweb', 'facebook', 'facebook', 'accuweather.com'],
    #                   9: ['https', 'https', 'https', 'https', 'https', 'https', 'https', 'https', 'https'],
    #                   10: ['https', 'https', 'https', 'https', 'instagram', 'instagram', 'unknown', 'unknown', 'unknown', 'facebook']}
#    print(f"Len {len(dictframel7app[i])}\n") # print length of l7app metadata for each frame

    dictframeodid[k] = odid[k]
    dictframedstaddr[k] = dstaddr[k].split(",")
    dictframedstport[k] = dstport[k].split(",")
# >>> dstport
# ["443','525','004','434','915','000','000','000','000','000','000','000','000','000','000','000','000'", '551,994,434,868,334,000,000,000,000,000,000,000,000,000,000,000', '36878,443,53855,58124,50800,25,55444,443,62415,443,54447,443', '59882,443,54318,443,62510,443,50355,80,57692', '562,134,434,086,539,000,000,000,000,000,000,000,000,000,000,000,000,000', '4,375,443,515,661,000,000,000,000,000,000,000,000,000,000,000,000', '443,51614,443,5007,443,64522,5222,52710,443,55344,80,59573,10050,49846', '533,464,435,944,144,000,000,000,000,000,000,000,000,000,000', '443,40807,443,49779,443,53295,443,38762,80', '4,434,195,744,352,680,000,000,000,000,000,000,000', '44,356,110,443,129,200,000,000,000,000,000,000,000,000']
# >>> dictframedstport
# {0: ["443'", "'525'", "'004'", "'434'", "'915'", "'000'", "'000'", "'000'", "'000'", "'000'", "'000'", "'000'", "'000'", "'000'", "'000'", "'000'", "'000'"], 1: ['551', '994', '434', '868', '334', '000', '000', '000', '000', '000', '000', '000', '000', '000', '000', '000'], 2: ['36878', '443', '53855', '58124', '50800', '25', '55444', '443', '62415', '443', '54447', '443'], 3: ['59882', '443', '54318', '443', '62510', '443', '50355', '80', '57692'], 4: ['562', '134', '434', '086', '539', '000', '000', '000', '000', '000', '000', '000', '000', '000', '000', '000', '000', '000'], 5: ['4', '375', '443', '515', '661', '000', '000', '000', '000', '000', '000', '000', '000', '000', '000', '000', '000'], 6: ['443', '51614', '443', '5007', '443', '64522', '5222', '52710', '443', '55344', '80', '59573', '10050', '49846'], 7: ['533', '464', '435', '944', '144', '000', '000', '000', '000', '000', '000', '000', '000', '000', '000'], 8: ['443', '40807', '443', '49779', '443', '53295', '443', '38762', '80'], 9: ['4', '434', '195', '744', '352', '680', '000', '000', '000', '000', '000', '000', '000'], 10: ['44', '356', '110', '443', '129', '200', '000', '000', '000', '000', '000', '000', '000', '000']}

#    dictframestarttime[k] = starttime[k].split(" GTB Daylight Time,")

# >>> starttime
# 'Jul 10, 2019 11:53:30.922000000 GTB Daylight Time,Jul 10, 2019 11:54:05.814000000 GTB Daylight Time,Jul 10, 2019 11:54:05.814000000 GTB Daylight Time,Jul 10, 2019 11:53:43.011000000 GTB Daylight Time,Jul 10, 2019 11:53:43.011000000 GTB Daylight Time,Jul 1'

for i in range(len(dictframel7app)):    # i = Excel column
    for ii in range(i+1,len(dictframel7app)):   # ii = next Excel columns
        if dictframeodid[i] != dictframeodid[ii]:   # !!! dictframeodid[i] != dictframeodid[ii] !!! only if ODID of frame x is different than one of frame y try to verify its l7app, its dstaddr and dstport for equality
            for j in range(len(dictframel7app[i])):         # length of current frame's l7app values
                for jj in range(len(dictframel7app[ii])):   # length of next frames' l7app values
                    if dictframedstaddr[i][j] == dictframedstaddr[ii][jj]:
                        #print(f"Destination address {dictframedstaddr[i][j]} in IPFIX ODID={int(dictframeodid[i])} frame {int(framenr[i])}, flow #{j+1} is the same with flow of frame {int(framenr[ii])}, flow #{jj+1} of ODID={int(dictframeodid[ii])}.\n")
                        if int(dictframedstport[i][j].replace("'", "")) == int(dictframedstport[ii][jj].replace("'", "")):  # the dictframedstport[i][j] contains str and is faster to compare integers
                            #print(f"The cflow.dstport=={dictframedstport[i][j]} in frame {int(framenr[i])}, flow #{j+1} is the same with frame {int(framenr[ii])}, flow #{jj+1}\n")
                            if dictframel7app[i][j] == dictframel7app[ii][jj]:
                                print(f"For debugging, at this time i={i}, j={j}, ii={ii}, jj={jj}")
                                flowstart1 = data[i]['_source']['layers']['cflow'][list((data[i]['_source']['layers']['cflow']).keys())[-1]][f"Flow {j+1} ({dictframel7app[i][j]})"]['cflow.timedelta_tree']['cflow.abstimestart'][0:28]     # refer to Why#2
                                flowstart2 = data[ii]['_source']['layers']['cflow'][list((data[ii]['_source']['layers']['cflow']).keys())[-1]][f"Flow {jj + 1} ({dictframel7app[i][j]})"]['cflow.timedelta_tree']['cflow.abstimestart'][0:28]
                                diff = msecdiff(flowstart1, flowstart2)
                                if diff <= 6000:    # only if IPFIX flowStartMilliseconds values difference is lower than 6 sec
                                    print(f"IPFIX ODID={int(dictframeodid[i])}, frame {int(framenr[i])}, flow #{j+1} <==> cflow.pie.ixia.l7-application-name=={dictframel7app[i][j]} && cflow.dstaddr=={dictframedstaddr[i][j]} && cflow.dport=={dictframedstport[i][j]} <==> ODID={int(dictframeodid[ii])}, frame {int(framenr[ii])}, flow #{jj+1}\n")
                                #commonflowsdoc.add_paragraph(f"IPFIX ODID={int(dictframeodid[i])}, frame {int(framenr[i])}, flow #{j+1} <==> cflow.pie.ixia.l7-application-name=={dictframel7app[i][j]} && cflow.dstaddr=={dictframedstaddr[i][j]} && cflow.dport=={dictframedstport[i][j]} <==> ODID={int(dictframeodid[ii])}, frame {int(framenr[ii])}, flow #{jj+1}\n")
                                    commonflowsdoc.add_paragraph(f"IPFIX ODID={int(dictframeodid[i])}, frame {int(framenr[i])}, flow #{j + 1} <==> cflow.pie.ixia.l7-application-name=={dictframel7app[i][j]} && cflow.dstaddr=={dictframedstaddr[i][j]} && cflow.dport=={dictframedstport[i][j]} <==> ODID={int(dictframeodid[ii])}, frame {int(framenr[ii])}, flow #{jj + 1} ==>>> The flowStartMilliseconds difference between these two is of {int(diff)} milliseconds!\n")
                                    commonflowsdoc.save(commonflowsdocAbsPath)

#----------------------------------------------------------------------------------------------------------------------------------------------
# @@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@ Why#1 @@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@
#
# >>> b
# 'Jul 10, 2019 11:53:30.922000000 GTB Daylight Time,Jul 10, 2019 11:54:05.814000000 GTB Daylight Time,Jul 10, 2019 11:54:05.814000000 GTB Daylight Time,Jul 10, 2019 11:53:43.011000000 GTB Daylight Time,Jul 10, 2019 11:53:43.011000000 GTB Daylight Time,Jul 1'
# >>> b.split(" GTB Daylight Time,")
# ['Jul 10, 2019 11:53:30.922000000', 'Jul 10, 2019 11:54:05.814000000', 'Jul 10, 2019 11:54:05.814000000', 'Jul 10, 2019 11:53:43.011000000', 'Jul 10, 2019 11:53:43.011000000', 'Jul 1']
#
# >>> date_time = "Jul 10, 2019 11:53:30.122000"
# >>> pattern = '%b %d, %Y %H:%M:%S.%f'
# >>> epoch1 = int(time.mktime(time.strptime(date_time, pattern)))	# conver current time to UNIX epoch time (seconds since 00:00:00 UTC on 1 January 1970)
# >>>
# >>> epoch1
# 1562748810
#
# >>> c = int( ('Jul 10, 2019 11:53:30.922000000'.split("."))[1])/10**6  # split returns a list of 2 elements, the 2nd one is after dot(.) so we get it by [1]
# >>>
# >>> epoch1*1000 + c
# 1562748810922.0
#
# def msecdiff(flow1, flow2):
#     return abs((int((flow1.split("."))[1]) / 10 ** 3 + int(time.mktime(time.strptime(flow1, pattern)))) - (int((flow2.split("."))[1]) / 10 ** 3 + int(time.mktime(time.strptime(flow2, pattern)))))  # diff in milliseconds
#
# msecdiff("Jul 10, 2019 11:54:15.210000","Jul 10, 2019 11:54:15.098000")
# 112.0
# msecdiff("Jul 10, 2019 11:52:08.558000","Jul 10, 2019 11:52:53.512000")
# 1.0
# msecdiff("Jul 10, 2019 11:53:09.569000","Jul 10, 2019 11:53:22.783000")
# 227.0
#----------------------------------------------------------------------------------------------------------------------------------------------
# @@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@ Why#2 @@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@
#
# >>> data[5]['_source']['layers']['cflow']['Set 1 [id=256] (12 flows)']['Flow 10 (https)']['cflow.timedelta_tree']['cflow.abstimestart']
# 'Jul 10, 2019 11:51:30.709000000 GTB Daylight Time'
#
# >>> data[5]['_source']['layers']['cflow']['Set 1 [id=256] (12 flows)']['Flow 10 (https)']['cflow.timedelta_tree']['cflow.abstimestart'][0:28]
# 'Jul 10, 2019 11:51:30.709000'
#
# >>> for j in range(0,10):	# j is the frame column index, the frame number if PCAP contains only NetFlow frames exported in JSON
#     print((data[j]['_source']['layers']['cflow'][list((data[j]['_source']['layers']['cflow']).keys())[-1]]).keys())
# with [list((data[i]['_source']['layers']['cflow']).keys())[-1]] we're accessing the last element in keys() of dict following "cflow", which is 'Set xx [id=256] (yy flows)' "
# dict_keys(['cflow.flowset_id', 'cflow.flowset_length', 'cflow.template_frame', 'Flow 1 (https)', 'Flow 2 (https)', 'Flow 3 (https)', 'Flow 4 (https)', 'Flow 5 (https)', 'Flow 6 (https)', 'Flow 7 (https)', 'Flow 8 (https)', 'Flow 9 (https)', 'Flow 10 (https)', 'Flow 11 (https)', 'Flow 12 (https)', 'Flow 13 (https)'])
# dict_keys(['cflow.flowset_id', 'cflow.flowset_length', 'cflow.template_frame', 'Flow 1 (chat)', 'Flow 2 (amazon-kindle)', 'Flow 3 (amazon-kindle)', 'Flow 4 (too_early)', 'Flow 5 (facebook)', 'Flow 6 (too_early)', 'Flow 7 (too_early)', 'Flow 8 (facebook)', 'Flow 9 (https)', 'Flow 10 (https)', 'Flow 11 (too_early)', 'Flow 12 (too_early)'])
# dict_keys(['cflow.flowset_id', 'cflow.flowset_length', 'cflow.template_frame', 'Flow 1 (too_early)', 'Flow 2 (too_early)', 'Flow 3 (too_early)', 'Flow 4 (too_early)', 'Flow 5 (too_early)', 'Flow 6 (smtp)', 'Flow 7 (smtp)', 'Flow 8 (office365)', 'Flow 9 (office365)', 'Flow 10 (https)', 'Flow 11 (https)', 'Flow 12 (facebook)'])
# dict_keys(['cflow.flowset_id', 'cflow.flowset_length', 'cflow.template_frame', 'Flow 1 (https)', 'Flow 2 (facebook)', 'Flow 3 (facebook)', 'Flow 4 (https)', 'Flow 5 (https)', 'Flow 6 (https)', 'Flow 7 (https)', 'Flow 8 (accuweather.com)', 'Flow 9 (accuweather.com)'])
# dict_keys(['cflow.flowset_id', 'cflow.flowset_length', 'cflow.template_frame', 'Flow 1 (facebook)', 'Flow 2 (https)', 'Flow 3 (https)', 'Flow 4 (too_early)', 'Flow 5 (facebook)', 'Flow 6 (facebook)', 'Flow 7 (too_early)', 'Flow 8 (too_early)', 'Flow 9 (https)', 'Flow 10 (https)', 'Flow 11 (hpvirtgrp)', 'Flow 12 (hpvirtgrp)', 'Flow 13 (https)'])
# dict_keys(['cflow.flowset_id', 'cflow.flowset_length', 'cflow.template_frame', 'Flow 1 (https)', 'Flow 2 (https)', 'Flow 3 (https)', 'Flow 4 (documentum_s)', 'Flow 5 (documentum_s)', 'Flow 6 (facebook)', 'Flow 7 (facebook)', 'Flow 8 (https)', 'Flow 9 (too_early)', 'Flow 10 (https)', 'Flow 11 (too_early)', 'Flow 12 (too_early)'])
# dict_keys(['cflow.flowset_id', 'cflow.flowset_length', 'cflow.template_frame', 'Flow 1 (https)', 'Flow 2 (https)', 'Flow 3 (https)', 'Flow 4 (https)', 'Flow 5 (youtube)', 'Flow 6 (youtube)', 'Flow 7 (chat)', 'Flow 8 (chat)', 'Flow 9 (https)', 'Flow 10 (https)', 'Flow 11 (http)', 'Flow 12 (http)', 'Flow 13 (zabbix-agent)', 'Flow 14 (zabbix-agent)'])
# dict_keys(['cflow.flowset_id', 'cflow.flowset_length', 'cflow.template_frame', 'Flow 1 (too_early)', 'Flow 2 (https)', 'Flow 3 (https)', 'Flow 4 (https)', 'Flow 5 (https)', 'Flow 6 (too_early)', 'Flow 7 (https)', 'Flow 8 (https)', 'Flow 9 (https)', 'Flow 10 (https)', 'Flow 11 (https)'])
# dict_keys(['cflow.flowset_id', 'cflow.flowset_length', 'cflow.template_frame', 'Flow 1 (youtube)', 'Flow 2 (youtube)', 'Flow 3 (hwwlcs.com)', 'Flow 4 (hwwlcs.com)', 'Flow 5 (icloudweb)', 'Flow 6 (icloudweb)', 'Flow 7 (facebook)', 'Flow 8 (facebook)', 'Flow 9 (accuweather.com)'])
# dict_keys(['cflow.flowset_id', 'cflow.flowset_length', 'cflow.template_frame', 'Flow 1 (https)', 'Flow 2 (https)', 'Flow 3 (https)', 'Flow 4 (https)', 'Flow 5 (https)', 'Flow 6 (https)', 'Flow 7 (https)', 'Flow 8 (https)', 'Flow 9 (https)'])

# How to debug:
# >> > x = 99
# >> > (data[x]['_source']['layers']['cflow'][list((data[x]['_source']['layers']['cflow']).keys())[-1]]).keys()
# dict_keys(['cflow.flowset_id', 'cflow.flowset_length', 'cflow.template_frame', 'Flow 1 (facebook)', 'Flow 2 (https)',
#            'Flow 3 (youtube)', 'Flow 4 (https)', 'Flow 5 (google)', 'Flow 6 (facebook)', 'Flow 7 (android)',
#            'Flow 8 (https)', 'Flow 9 (https)', 'Flow 10 (chat)'])
#
# >> > listi = ['Flow 1 (facebook)', 'Flow 2 (https)', 'Flow 3 (youtube)', 'Flow 4 (https)', 'Flow 5 (google)',
#               'Flow 6 (facebook)', 'Flow 7 (android)', 'Flow 8 (https)', 'Flow 9 (https)', 'Flow 10 (chat)']
#
# >> > for i in listi:
#     print((data[x]['_source']['layers']['cflow'][list((data[x]['_source']['layers']['cflow']).keys())[-1]])[i][
#               'cflow.timedelta_tree']['cflow.abstimestart'][0:28])
#
# May
# 19, 2020
# 10: 37:35.899000
# :
# May
# 19, 2020
# 10: 37:19.876000

#----------------------------------------------------------------------------------------------------------------------------------------------
# @@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@ Why#3 @@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@
#
# for Export AppStack for IXIA 2 <=>  NF_L7app_Dport_Daddr.xlsx <=> sh = xlrd.open_workbook("C:\\Users\\boburciu\\Desktop\\NF_L7app_Dport_Daddr.xlsx").sheet_by_index(0)
# >>>l7app = sh.col_values(11, start_rowx=0)
# >>>l7app
# ['L7 Application Name', 'https,https,https,https,https,https,https,https,https,https,https,https,https', 'chat,amazon-kindle,amazon-kindle,too_early,facebook,too_early,too_early,facebook,https,https,too_early,too_early', 'too_early,too_early,too_early,too_early,too_early,smtp,smtp,office365,office365,https,https,facebook', 'https,facebook,facebook,https,https,https,https,accuweather.com,accuweather.com', 'facebook,https,https,too_early,facebook,facebook,too_early,too_early,https,https,hpvirtgrp,hpvirtgrp,https', 'https,https,https,documentum_s,documentum_s,facebook,facebook,https,too_early,https,too_early,too_early', 'https,https,https,https,youtube,youtube,chat,chat,https,https,http,http,zabbix-agent,zabbix-agent', 'too_early,https,https,https,https,too_early,https,https,https,https,https', 'youtube,youtube,hwwlcs.com,hwwlcs.com,icloudweb,icloudweb,facebook,facebook,accuweather.com', 'https,https,https,https,https,https,https,https,https', 'https,https,https,https,instagram,instagram,unknown,unknown,unknown,facebook']
# >>>l7app[1]
# 'https,https,https,https,https,https,https,https,https,https,https,https,https'
# >>>l7app[1].split(",")
# ['https', 'https', 'https', 'https', 'https', 'https', 'https', 'https', 'https', 'https', 'https', 'https', 'https']
